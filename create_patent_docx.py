from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Patent / Design Application Details for AEGIS-1', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Inventors and Applicants
doc.add_heading('1. Full Name and Address of the inventor(s) and applicant(s)', level=1)
doc.add_paragraph('Inventor(s): [Insert Your Name / Team Members\' Names Here]')
doc.add_paragraph('Applicant(s): [Insert Applicant/Institute Name Here]')
doc.add_paragraph('Address: [Insert Full Postal Address Here]')

# 2. Suggested Title
doc.add_heading('2. Suggested title of the Design', level=1)
doc.add_paragraph('AEGIS-1: Hybrid Deep Learning Framework and Architecture for Real-Time Tactical Drone Surveillance')

# 3. Information about Design
doc.add_heading('3. Information about your design', level=1)

doc.add_heading('Introduction:', level=2)
doc.add_paragraph('AEGIS-1 is a cost-effective, hybrid deep-learning framework designed for real-time military surveillance using unmanned aerial vehicles (UAVs). It addresses the challenge of deploying computationally intensive AI on resource-constrained platforms by utilizing a decoupled edge-cloud architecture, combining a lightweight hardware sensor node with a sophisticated multithreaded inference server.')

doc.add_heading('Working:', level=2)
doc.add_paragraph('The system utilizes an ESP32-CAM module as an edge device to acquire and stream MJPEG video data. This stream is ingested by a centralized backend processing server. The core working principle revolves around a highly parallelized multi-threaded architecture that runs a dual-model inference engine. It employs a pre-trained YOLOv8 model for rapid, generalized object localization and a custom-trained TensorFlow Lite (TFLite) "MIL-SCAN" model running concurrently to classify specific military hardware (helicopters, planes, tanks, vehicles, missiles) with high precision.')

doc.add_heading('Advantages:', level=2)
doc.add_paragraph('- Cost-Effective: Leverages low-cost edge sensors (ESP32-CAM) while offloading heavy computation.')
doc.add_paragraph('- Low Latency & High FPS: Multithreaded architecture with frame-skipping and strict lock synchronization eliminates sequential processing bottlenecks.')
doc.add_paragraph('- High Accuracy: The hybrid approach ensures both general spatial awareness and precise high-value target classification.')
doc.add_paragraph('- Dual-Mode Capability: Features both a "Live" streaming mode and an "Analysis" offline mode for high-resolution deep-scan intelligence gathering.')

doc.add_heading('Innovations and New Features:', level=2)
doc.add_paragraph('1. Hybrid Dual-Model AI Pipeline: Simultaneous execution of YOLOv8 and TFLite models on the same stream.')
doc.add_paragraph('2. Decoupled Processing: Separation of stream ingestion from AI execution threads using shared memory locks.')
doc.add_paragraph('3. strict deduplication logic and visual hierarchical rendering (YOLO in green, MIL-SCAN in red) for the operator HUD.')

doc.add_heading('Claims:', level=2)
doc.add_paragraph('1. A tactical surveillance system architecture comprising a lightweight edge-streaming node and a centralized multithreaded inference server.')
doc.add_paragraph('2. A hybrid object detection method that simultaneously applies generalized (YOLO) and specialized (TFLite) neural networks to a single live video feed.')
doc.add_paragraph('3. A dynamic operating mode switching mechanism allowing real-time AI inference and offline high-resolution deep-scan analysis within the same framework.')

# 4. Photographs/Drawings
doc.add_heading('4. Photographs/Drawings of the article in SEVEN VIEWS', level=1)
doc.add_paragraph('[INSTRUCTION: Please insert line diagrams, photographs, or CAD drawings of the AEGIS-1 hardware setup (drone + ESP32-CAM) or the system architecture in the following views]')
doc.add_paragraph('- Front View\n- Rear View\n- Top Plan View\n- Bottom Plan View\n- Left Side Elevation\n- Right Side Elevation\n- Isometric View')

# 5. Conclusion
doc.add_heading('5. Conclusion', level=1)
doc.add_paragraph('The AEGIS-1 system demonstrates a highly efficient, scalable, and innovative approach to tactical drone surveillance. By intelligently combining low-cost edge hardware with a powerful, multithreaded AI inference engine, it provides real-time, accurate threat intelligence without the prohibitive costs or hardware limitations typical of conventional military systems.')

doc.save('AEGIS-1_Patent_Draft.docx')
print("Patent draft saved to AEGIS-1_Patent_Draft.docx")
